# build_4p2_docx.py
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches

ROOT = Path(".")  # /path/to/EEG_48sounds で実行する前提
OUT_DOCX = ROOT / "渡邉俊介＿卒業論文_4.2_最強完全版.docx"

FIG_DIR = ROOT / "moduleB_outputs" / "figures"
TAB_DIR = ROOT / "moduleB_outputs" / "tables"

TASKS = [
    "emo_arousal","emo_approach","emo_valence",
    "emo_arousal_high","emo_approach_high","emo_valence_high",
    "is_ambiguous","category_3"
]

def add_caption(doc, caption):
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True

def add_fig(doc, fig_path: Path, caption: str, width_in=6.3):
    if not fig_path.exists():
        doc.add_paragraph(f"[図が見つかりません] {fig_path}")
        return
    doc.add_picture(str(fig_path), width=Inches(width_in))
    add_caption(doc, caption)

def load_summary_all():
    rows = []
    for t in TASKS:
        p = TAB_DIR / f"moduleB_summary_{t}_linear.csv"
        df = pd.read_csv(p)
        rows.append(df.iloc[0].to_dict())
    return pd.DataFrame(rows)

def main():
    doc = Document()

    # Title
    doc.add_heading("4.2 EEG単独で主観ターゲット・カテゴリをどこまで推定できるか", level=1)

    # --- 本文（ここは上で渡した「4.2 完全版本文」を貼るのが確実）
    doc.add_paragraph("【ここに、私が上で提示した4.2本文をそのまま貼ってください】")

    # --- 表4-2-1
    doc.add_heading("表4-2-1：EEG単独のLOSO性能サマリ（8ターゲット）", level=2)
    df = load_summary_all()

    # 表に載せたい列（環境で列名が微妙に違う可能性があるので、存在するものだけ採用）
    preferred_cols = ["task","kind","score_mean","score_sd","p_perm","n_subjects","n_trials"]
    cols = [c for c in preferred_cols if c in df.columns]
    show = df[cols].copy()

    table = doc.add_table(rows=1, cols=len(cols))
    hdr = table.rows[0].cells
    for j,c in enumerate(cols):
        hdr[j].text = c

    for _,r in show.iterrows():
        cells = table.add_row().cells
        for j,c in enumerate(cols):
            cells[j].text = str(r[c])

    # --- 図4-2-1〜8（LOSO）
    doc.add_heading("図4-2-1〜8：EEG単独のLOSO性能（ターゲット別）", level=2)
    for i,t in enumerate(TASKS, start=1):
        fig = FIG_DIR / f"FIG1_LOSO_{t}_linear.png"
        add_fig(doc, fig, f"図4-2-{i}：LOSO性能（{t}）", width_in=6.6)

    # --- MAP（ERP/TFR）
    doc.add_heading("図4-2-9〜：時間窓別推定性能（MAP）", level=2)
    map_dir = FIG_DIR / "maps_0to5000_v1"
    k = 9
    for t in TASKS:
        add_fig(doc, map_dir / f"MAP_ERP_{t}_score_line.png", f"図4-2-{k}：ERP時間窓別性能（{t}）"); k += 1
    for t in TASKS:
        add_fig(doc, map_dir / f"MAP_TFR_{t}_score_heatmap.png", f"図4-2-{k}：TFR帯域×時間窓 性能（{t}）"); k += 1

    # --- 重要度（分解＋詳細）
    imp_dir = FIG_DIR / "importance_decompose_0to5000_v1"
    doc.add_heading("重要度（ERP vs TFR短期 vs TFR長期、および詳細）", level=2)

    for t in TASKS:
        add_fig(doc, imp_dir / f"Fig_imp_group_{t}.png", f"図4-2-{k}：重要度分解（{t}）"); k += 1
    for t in TASKS:
        add_fig(doc, imp_dir / f"Fig_imp_ERP_ch_time_{t}.png", f"図4-2-{k}：ERP重要度（ch×time）（{t}）"); k += 1
    for t in TASKS:
        add_fig(doc, imp_dir / f"Fig_imp_TFRshort_band_time_{t}.png", f"図4-2-{k}：TFR短期重要度（band×time）（{t}）"); k += 1
    for t in TASKS:
        add_fig(doc, imp_dir / f"Fig_imp_TFRlong_band_time_{t}.png", f"図4-2-{k}：TFR長期重要度（band×time）（{t}）"); k += 1

    # --- ROI LOSO（あれば全部入れる：フォルダ内PNGを全挿入）
    roi_dir = FIG_DIR / "roi_loso_0to5000_v1"
    if roi_dir.exists():
        doc.add_heading("ROI別LOSO（領域限定での汎化）", level=2)
        for p in sorted(roi_dir.glob("*.png")):
            add_fig(doc, p, f"図4-2-{k}：ROI別LOSO（{p.name}）", width_in=6.6); k += 1

    doc.save(str(OUT_DOCX))
    print("saved:", OUT_DOCX)

if __name__ == "__main__":
    main()
